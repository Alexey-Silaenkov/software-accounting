import docx
import os
from datetime import date


class ImportDoc:


    # Создание и добавление стандартных параметров документа
    doc = docx.Document()
    doc.add_picture('modules/images/Ychpo.jpg', width = docx.shared.Cm(10))
    doc.add_heading(f'Отчет по выданным лицензионным ключам на {date.today()} \n', 1)
    
    # добавление параграфов в документ
    def importDoc(self, user: str, po: str, kluch: str):
        # добавляем параграф
        self.doc.add_paragraph(f'Пользователь: {user} \t ПО: {po} \t Ключ: {kluch[:-5]}')

        
    # Сохранение документа
    def saveDoc(self):
        user = os.environ.get( "USERNAME" )
        self.doc.save(f'C:\\users\{user}\Desktop\{date.today()} отчет.docx')


# a = ImportDoc()
# a.importDoc("Иванов Иван Иванович","Adobe V3","GHBR-UAB6-LSUF-87SD")
# a.importDoc("Иванов Иван Иванович","Adobe V3","GHBR-UAB6-LSUF-87SD")
# a.importDoc("Иванов Иван Иванович","Adobe V3","GHBR-UAB6-LSUF-87SD")
# a.importDoc("Иванов Иван Иванович","Adobe V3","GHBR-UAB6-LSUF-87SD")
# a.importDoc("Иванов Иван Иванович","Adobe V3","GHBR-UAB6-LSUF-87SD")
# a.importDoc("Иванов Иван Иванович","Adobe V3","GHBR-UAB6-LSUF-87SD")
# a.importDoc("Иванов Иван Иванович","Adobe V3","GHBR-UAB6-LSUF-87SD")
# a.importDoc("Иванов Иван Иванович","Adobe V3","GHBR-UAB6-LSUF-87SD")
# a.saveDoc()